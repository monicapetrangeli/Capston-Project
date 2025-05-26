import os
import time
import json
import fitz  # PyMuPDF
import re
import urllib3
import requests
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fuzzywuzzy import fuzz
from colorama import Fore, Style, init

# Initialize colorama for cross-platform colored terminal output
init()

# --- SETTINGS ---
FOLDER_PATH = "General surgery Ana Garcia Lopez"
OUTPUT_FOLDER = "Extracted_Documents"
FINAL_DOCX = "Final_Extracted_Report.docx"
TEMPLATE_DOCX = "Template.docx"
FILLED_TEMPLATE = "Filled_Template.docx"
CONFIDENCE_THRESHOLD = 0.6  # Minimum confidence to accept a field without human review
FIELD_STATS_FILE = "field_extraction_stats.csv"
API_TOKEN = "YOUR_ACCESS_TOKEN"  # Replace with your API token

# --- FIELD DEFINITIONS ---
FIELDS = [
    "DATE", "INCIDENT_DATE", "PATIENT_NAME", "DOB", "MEDICAL_DOCTOR", "LAWYER",
    "SPECIALITY", "CENTER", "JUDGE", "REFERENCE", "SUMMARY", "PRACTICE_DETAILS",
    "FORENSIC_REPORT", "EVALUATION_UML", "DIAGNOSTIC_CODE", "PROCEDURE_CODE",
    "SEQUELS", "DEFINITIVE?", "QUANTUM", "RISK_ASSESSMENT"
]
BATCH_SIZE = 3

# --- FIELD CATEGORIES FOR CONFIDENCE SCORING ---
PERSON_FIELDS = ["PATIENT_NAME", "MEDICAL_DOCTOR", "LAWYER", "JUDGE"]
DATE_FIELDS = ["DATE", "INCIDENT_DATE", "DOB"]
CODE_FIELDS = ["DIAGNOSTIC_CODE", "PROCEDURE_CODE"]
LONG_TEXT_FIELDS = ["SUMMARY", "PRACTICE_DETAILS", "FORENSIC_REPORT"]

# --- FUNCTION TO CALL THE API ---
def call_model(prompt):
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    destination_url = "https://grupmedai-api-des.itcomb.cat/llm/generate"
    headers = {
        'Authorization':  f'Bearer {API_TOKEN}',
        'Content-Type': 'application/json'
    }
    data = {
        "model": "llama3.1:8b",
        "prompt": prompt 
    }
    try:
        response = requests.post(destination_url, headers=headers, data=json.dumps(data), verify=False)
        if response.status_code == 200:
            response_data = response.json().get('text', '')
            return response_data
        else:
            print(f"API error: Status code {response.status_code}")
            return False
    except Exception as e:
        print(f"Request error: {e}")
        return False

# --- FUNCTION TO BATCH FIELDS ---
def batch_fields(field_list, batch_size):
    for i in range(0, len(field_list), batch_size):
        yield field_list[i:i+batch_size]

# --- FUNCTION TO READ PDFs ---
def read_pdfs_from_folder(folder_path):
    documents = []
    filenames = []
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".pdf"):
            file_path = os.path.join(folder_path, filename)
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            documents.append(text)
            filenames.append(os.path.splitext(filename)[0])
    return documents, filenames

# --- FUNCTION TO BUILD PROMPT WITH CONFIDENCE SCORING ---
def build_prompt(document_text, fields_batch, previous_values=None):
    fields_prompt = "\n".join([f"- {field}" for field in fields_batch])
    
    previous_context = ""
    if previous_values:
        previous_context = "Previously extracted fields:\n"
        for field, data in previous_values.items():
            if field not in fields_batch:  # Only include fields not in the current batch
                previous_context += f"- {field}: {data['value']}\n"
    
    prompt = f"""
You are an AI assistant extracting key information from a medico-legal report.

Instructions:
- For each field, extract only the direct, specific answer.
- For PERSON fields (like PATIENT_NAME, MEDICAL_DOCTOR), extract only the full name. No age, specialty, or biography.
- For DATE fields, extract only the date in format DD/MM/YYYY.
- Be concise. No paragraphs, no explanations, no extra text.
- If the field is missing in the document, respond with "Not Found".
- For each field, also provide a confidence score between 0.0 and 1.0, where:
  * 1.0 means you are absolutely certain the information is correct
  * 0.0 means you are completely uncertain or the information is not present
  * Values from 0.6 to 0.8 indicate moderate confidence
- YOU MUST RESPOND WITH VALID JSON FORMAT ONLY.
- Format your response as: {{"FIELD1": {{"value": "extracted value", "confidence": 0.95}}, "FIELD2": {{"value": "extracted value", "confidence": 0.7}}}}
- Do not include any additional text outside of the JSON object.

{previous_context}

Fields to extract:
{fields_prompt}

Document Text:
{document_text}
"""
    return prompt

# --- FUNCTION TO BUILD REFINEMENT PROMPT ---
def build_refinement_prompt(document_text, field, previous_value, user_feedback):
    prompt = f"""
You are an AI assistant extracting specific information from a medico-legal report.

You previously extracted the following information:
- Field: {field}
- Extracted value: "{previous_value['value']}"
- Your confidence: {previous_value['confidence']}

The user has provided feedback that this extraction needs correction.
User feedback: "{user_feedback}"

Instructions:
- Re-analyze the document text and correct your extraction for the field {field}.
- Only extract the specific value for this field, following these guidelines:
  * For PERSON fields: Extract only the full name.
  * For DATE fields: Extract only the date in format DD/MM/YYYY.
  * Be concise. No explanations or extra text.
- If the field is truly missing in the document, respond with "Not Found".
- Provide a new confidence score between 0.0 and 1.0.
- YOU MUST RESPOND WITH VALID JSON FORMAT ONLY.
- Format: {{"value": "corrected extraction", "confidence": 0.95}}
- Do not include any additional text outside of the JSON object.

Document Text:
{document_text}
"""
    return prompt

# --- FUNCTION TO QUERY LLM ---
def query_llm(prompt):
    try:
        response = call_model(prompt)
        return response
    except Exception as e:
        print(f"Error querying LLM: {e}")
        return None

# --- FUNCTION TO CLEAN LLM RESPONSE ---
def clean_json_response(response_text):
    if response_text is None or response_text is False:
        return None
        
    # Find JSON object in response if it's embedded in other text
    json_pattern = r'({.*})'
    match = re.search(json_pattern, response_text, re.DOTALL)
    
    if match:
        potential_json = match.group(1)
        # Clean up common JSON formatting issues
        potential_json = potential_json.replace("'", '"')  # Replace single quotes with double quotes
        potential_json = re.sub(r',\s*}', '}', potential_json)  # Remove trailing commas
        potential_json = re.sub(r'//.*?\n', '\n', potential_json)  # Remove single-line comments
        potential_json = re.sub(r'/\*.*?\*/', '', potential_json, flags=re.DOTALL)  # Remove multi-line comments
        
        # Fix unquoted "Not Found" values - this is the main fix for your error
        potential_json = re.sub(r':\s*Not Found\s*,', ': "Not Found",', potential_json)
        potential_json = re.sub(r':\s*Not Found\s*}', ': "Not Found"}', potential_json)
        
        # Fix other common unquoted values
        potential_json = re.sub(r':\s*([A-Za-z][A-Za-z0-9\s]*)\s*,', r': "\1",', potential_json)
        potential_json = re.sub(r':\s*([A-Za-z][A-Za-z0-9\s]*)\s*}', r': "\1"}', potential_json)
        
        return potential_json
    else:
        return response_text  # Return original if no JSON-like pattern found

# --- FUNCTION TO PARSE RESPONSE WITH CONFIDENCE ---
def parse_response(response_text, fields_batch):
    field_values = {}
    
    # Initialize with "Not Found" values and zero confidence
    for field in fields_batch:
        field_values[field] = {"value": "Not Found", "confidence": 0.0}
    
    if response_text is None or response_text is False:
        return field_values
    
    # Clean the response to get valid JSON
    cleaned_response = clean_json_response(response_text)
    
    try:
        data = json.loads(cleaned_response)
        for field in fields_batch:
            if field in data:
                field_data = data[field]
                # Handle if return is already in the expected format
                if isinstance(field_data, dict) and "value" in field_data:
                    value = field_data.get("value", "Not Found")
                    confidence = field_data.get("confidence", 0.0)
                # Handle if return is just a string value
                elif isinstance(field_data, str):
                    value = field_data
                    confidence = 0.5  # Default confidence for legacy format
                # Handle if return has a nested 'Name' field
                elif isinstance(field_data, dict) and "Name" in field_data:
                    value = field_data["Name"]
                    confidence = 0.5  # Default confidence for legacy format
                else:
                    value = "Not Found"
                    confidence = 0.0
                    
                field_values[field] = {"value": value, "confidence": confidence}
    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
        print(f"Response text: {response_text}")
        print(f"Cleaned response: {cleaned_response}")
        
        # Fallback: Try to extract field-value pairs directly from text using regex
        for field in fields_batch:
            # Pattern for the new format with confidence
            new_pattern = rf'"{field}"\s*:\s*{{\s*"value"\s*:\s*"([^"]*)"\s*,\s*"confidence"\s*:\s*([\d.]+)'
            new_match = re.search(new_pattern, cleaned_response)
            
            if new_match:
                value = new_match.group(1)
                confidence = float(new_match.group(2))
                field_values[field] = {"value": value, "confidence": confidence}
            else:
                # Legacy pattern for just the value
                old_pattern = rf'"{field}"\s*:\s*"([^"]*)"'
                old_match = re.search(old_pattern, cleaned_response)
                if old_match:
                    value = old_match.group(1)
                    field_values[field] = {"value": value, "confidence": 0.5}  # Default confidence
    
    return field_values

# --- FUNCTION TO CALCULATE CONFIDENCE BASED ON DOCUMENT TEXT ---
def calculate_external_confidence(field, extracted_value, document_text):
    """Calculate an additional confidence score based on text matching"""
    if extracted_value == "Not Found" or not extracted_value:
        return 0.0
        
    # Different confidence calculation strategies based on field type
    if field in PERSON_FIELDS:
        # For person fields, check if name appears in the document
        confidence = 0.0
        if extracted_value in document_text:
            confidence = 0.9  # Direct match
        else:
            # Split the name and check for partial matches
            name_parts = extracted_value.split()
            if len(name_parts) > 1:
                matches = sum(1 for part in name_parts if part in document_text)
                confidence = min(0.7, matches / len(name_parts))
                
            # Use fuzzy matching as a backup
            highest_ratio = 0
            for line in document_text.split('\n'):
                ratio = fuzz.partial_ratio(extracted_value.lower(), line.lower())
                highest_ratio = max(highest_ratio, ratio)
            confidence = max(confidence, highest_ratio / 100)
                
        return confidence
        
    elif field in DATE_FIELDS:
        # For date fields, check against date patterns
        date_patterns = [
            r'\d{1,2}/\d{1,2}/\d{2,4}',  # DD/MM/YYYY or MM/DD/YYYY
            r'\d{1,2}-\d{1,2}-\d{2,4}',  # DD-MM-YYYY
            r'\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4}',  # DD Mon YYYY
        ]
        
        # Check if the extracted date follows standard patterns
        date_format_match = any(re.match(pattern, extracted_value) for pattern in date_patterns)
        
        # Check if the date appears in the document
        date_in_doc = extracted_value in document_text
        
        # Calculate confidence based on format and presence
        if date_format_match and date_in_doc:
            return 0.9
        elif date_format_match:
            return 0.7
        elif date_in_doc:
            return 0.6
        else:
            # Use fuzzy matching as a backup
            highest_ratio = 0
            for line in document_text.split('\n'):
                ratio = fuzz.partial_ratio(extracted_value, line)
                highest_ratio = max(highest_ratio, ratio)
            return highest_ratio / 100
            
    elif field in CODE_FIELDS:
        # For medical/diagnostic codes, check for format and presence
        code_in_doc = extracted_value in document_text
        if code_in_doc:
            return 0.9
        else:
            # Use fuzzy matching for codes that might have different formatting
            highest_ratio = 0
            for line in document_text.split('\n'):
                ratio = fuzz.partial_ratio(extracted_value, line)
                highest_ratio = max(highest_ratio, ratio)
            return highest_ratio / 100
            
    elif field in LONG_TEXT_FIELDS:
        # For long text fields, check if key phrases appear in the document
        if len(extracted_value) < 10:  # Very short summary is suspicious
            return 0.3
        
        # Check for semantic similarity with document sections
        words = extracted_value.lower().split()
        word_count = len(words)
        
        # Count how many words from the extraction appear in the document
        word_matches = sum(1 for word in words if word.lower() in document_text.lower())
        
        if word_count > 0:
            word_match_ratio = word_matches / word_count
            return min(0.8, word_match_ratio)  # Cap at 0.8 for text summaries
        else:
            return 0.0
    
    else:
        # For other fields, use basic fuzzy matching
        highest_ratio = 0
        for line in document_text.split('\n'):
            ratio = fuzz.partial_ratio(extracted_value.lower(), line.lower())
            highest_ratio = max(highest_ratio, ratio)
        return highest_ratio / 100

# --- FUNCTION TO COMBINE CONFIDENCES ---
def combine_confidences(llm_confidence, external_confidence):
    """Combine the LLM's confidence with our external validation"""
    # Weight the LLM confidence higher than our basic validation
    combined = (0.7 * llm_confidence) + (0.3 * external_confidence)
    return min(1.0, combined)  # Cap at 1.0

# --- FUNCTION TO CREATE FINAL DOCX ---
def create_final_docx(data, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for field in FIELDS:
        field_data = data.get(field, {"value": "Not Found", "confidence": 0.0})
        value = field_data["value"]
        confidence = field_data["confidence"]
        
        # Format confidence as a percentage
        confidence_str = f"{confidence:.1%}"
        
        # Convert any non-string values to strings
        if not isinstance(value, str):
            value = str(value)
            
        # Add a paragraph with both value and confidence
        doc.add_paragraph(f"{field}: {value} (Confidence: {confidence_str})")

    doc.save(output_path)

# --- FUNCTION TO ENSURE STRING VALUES ---
def ensure_string_values(data_dict):
    """Convert all values in a dictionary to strings to prevent template filling errors."""
    processed_dict = {}
    for key, value in data_dict.items():
        if isinstance(value, dict) and "value" in value:
            processed_dict[key] = str(value["value"])
        elif not isinstance(value, str):
            processed_dict[key] = str(value)
        else:
            processed_dict[key] = value
    return processed_dict

# --- FUNCTION TO FILL TEMPLATE ---
def fill_template(data, template_path, output_path):
    try:
        # Convert data format from {field: {value, confidence}} to {field: value}
        template_data = {}
        for field, field_data in data.items():
            if isinstance(field_data, dict) and "value" in field_data:
                template_data[field] = field_data["value"]
            else:
                template_data[field] = field_data
                
        # Convert all data values to strings to prevent replace() errors
        template_data = ensure_string_values(template_data)
        
        # Load the template document
        doc = Document(template_path)
        
        # Process each table in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Check each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        # Find all placeholders in the format {FIELD_NAME}
                        placeholder_pattern = r'\{([A-Z_?]+)\}'
                        matches = re.findall(placeholder_pattern, paragraph.text)
                        
                        # Replace placeholders with extracted values
                        for field in matches:
                            if field in template_data:
                                value = template_data.get(field, "Not Found")
                                paragraph.text = paragraph.text.replace(f"{{{field}}}", value)
        
        # Process paragraphs outside of tables
        for paragraph in doc.paragraphs:
            placeholder_pattern = r'\{([A-Z_?]+)\}'
            matches = re.findall(placeholder_pattern, paragraph.text)
            
            for field in matches:
                if field in template_data:
                    value = template_data.get(field, "Not Found")
                    paragraph.text = paragraph.text.replace(f"{{{field}}}", value)
        
        # Save the filled template
        doc.save(output_path)
        print(f"Template filled and saved to: {output_path}")
        return True
    except Exception as e:
        print(f"Error filling template: {e}")
        return False

# --- IMPROVED FUNCTION TO FILL TEMPLATE ---
def fill_template_improved(data, template_path, output_path):
    """
    Improved template filling that better handles complex documents and run-level formatting.
    """
    try:
        # Convert data format from {field: {value, confidence}} to {field: value}
        template_data = {}
        for field, field_data in data.items():
            if isinstance(field_data, dict) and "value" in field_data:
                template_data[field] = field_data["value"]
            else:
                template_data[field] = field_data
                
        # Convert all data values to strings
        template_data = ensure_string_values(template_data)
        
        # Load the template document
        doc = Document(template_path)
        
        # Helper function to replace text while preserving formatting
        def replace_placeholder_in_paragraph(paragraph, field, value):
            # Get placeholder text pattern
            placeholder = f"{{{field}}}"
            
            # Skip if placeholder not in paragraph
            if placeholder not in paragraph.text:
                return False
                
            # Identify which run contains the placeholder
            found = False
            runs_with_placeholder = []
            placeholder_start_run = None
            placeholder_text_found = ""
            
            # First find all runs containing parts of the placeholder
            for i, run in enumerate(paragraph.runs):
                if placeholder in run.text:
                    # Simple case: entire placeholder in one run
                    run.text = run.text.replace(placeholder, value)
                    found = True
                    break
                    
                # Check for partial placeholder (placeholder split across runs)
                if placeholder_start_run is None and "{" in run.text and run.text.endswith("{"):
                    # Start of a potential placeholder
                    placeholder_start_run = i
                    placeholder_text_found = run.text[run.text.index("{"):]
                elif placeholder_start_run is not None:
                    # Continue building the placeholder
                    placeholder_text_found += run.text
                    
                    # Check if we've found the complete placeholder
                    if placeholder in placeholder_text_found:
                        runs_with_placeholder = list(range(placeholder_start_run, i + 1))
                        break
                    
                    # Check if this run closes a placeholder but not our target
                    if "}" in run.text:
                        placeholder_start_run = None
                        placeholder_text_found = ""
            
            # Handle multi-run placeholders
            if not found and runs_with_placeholder:
                # First run gets everything up to the placeholder start + the replacement value
                start_run = paragraph.runs[runs_with_placeholder[0]]
                placeholder_start_pos = start_run.text.find("{")
                start_run.text = start_run.text[:placeholder_start_pos] + value
                
                # Delete text from middle runs
                for i in runs_with_placeholder[1:-1]:
                    paragraph.runs[i].text = ""
                
                # Last run gets everything after the placeholder end
                end_run = paragraph.runs[runs_with_placeholder[-1]]
                placeholder_end_pos = end_run.text.find("}") + 1
                if placeholder_end_pos < len(end_run.text):
                    end_run.text = end_run.text[placeholder_end_pos:]
                else:
                    end_run.text = ""
                    
                found = True
                
            return found
        
        # Process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for field, value in template_data.items():
                            replace_placeholder_in_paragraph(paragraph, field, value)
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            for field, value in template_data.items():
                replace_placeholder_in_paragraph(paragraph, field, value)
                
        # Save the filled template
        doc.save(output_path)
        print(f"Template filled and saved to: {output_path}")
        return True
    except Exception as e:
        print(f"Error filling template (improved method): {e}")
        return False

# --- FUNCTION TO PRINT EXTRACTION RESULTS ---
def print_extraction_results(result):
    print("\nExtraction Results:")
    print("-" * 60)
    for field, data in result.items():
        value = data["value"]
        confidence = data["confidence"]
        
        # Color-code based on confidence
        if confidence >= 0.8:
            confidence_color = Fore.GREEN
        elif confidence >= CONFIDENCE_THRESHOLD:
            confidence_color = Fore.YELLOW
        else:
            confidence_color = Fore.RED
            
        confidence_str = f"{confidence:.1%}"
        
        # Show field name and value, or "Not Found" if empty/missing
        if value and value != "Not Found" and value.strip():
            print(f"{field}: {value} {confidence_color}(Confidence: {confidence_str}){Style.RESET_ALL}")
        else:
            print(f"{field}: Not Found")
    print("-" * 60)

# --- FUNCTION TO INTERACT WITH USER FOR FEEDBACK ---
def get_user_feedback(extraction_results, document_text, confidence_threshold=CONFIDENCE_THRESHOLD):
    """Interactive loop to get user feedback on extraction results"""
    print("\n" + "="*80)
    print(f"{Fore.CYAN}EXTRACTION REVIEW{Style.RESET_ALL}")
    print("="*80)
    print(f"Fields with confidence below {confidence_threshold:.0%} will be flagged for review.")
    print("You can accept, edit, or request re-extraction for each field.")
    print("-"*80)
    
    updated_results = extraction_results.copy()
    fields_to_review = []
    
    # First, identify all fields that need review
    for field, data in extraction_results.items():
        confidence = data["confidence"]
        value = data["value"]
        
        # Flag fields with low confidence or "Not Found" values for review
        if confidence < confidence_threshold or value == "Not Found":
            fields_to_review.append(field)
    
    # If nothing to review, return original results
    if not fields_to_review:
        print(f"{Fore.GREEN}All fields extracted with high confidence. No review needed.{Style.RESET_ALL}")
        return updated_results
    
    # Process each field that needs review
    for field in fields_to_review:
        data = extraction_results[field]
        confidence = data["confidence"]
        value = data["value"]
        confidence_str = f"{confidence:.1%}"
        
        # Color code based on confidence
        if confidence >= 0.8:
            confidence_color = Fore.GREEN
        elif confidence >= confidence_threshold:
            confidence_color = Fore.YELLOW
        else:
            confidence_color = Fore.RED
            
        print(f"\nReviewing field: {Fore.CYAN}{field}{Style.RESET_ALL}")
        print(f"Extracted value: \"{value}\"")
        print(f"Confidence: {confidence_color}{confidence_str}{Style.RESET_ALL}")
        
        print("\nOptions:")
        print("1. Accept current value")
        print("2. Edit value manually")
        print("3. Request re-extraction with feedback")
        print("4. Skip for now")
        
        choice = input("\nYour choice (1-4): ").strip()
        
        if choice == "1":
            print(f"{Fore.GREEN}Value accepted.{Style.RESET_ALL}")
            # Keep existing value but adjust confidence to reflect human verification
            updated_results[field]["confidence"] = 1.0
            
        elif choice == "2":
            new_value = input(f"Enter new value for {field}: ").strip()
            updated_results[field]["value"] = new_value
            updated_results[field]["confidence"] = 1.0  # Human-provided value gets 100% confidence
            print(f"{Fore.GREEN}Value updated.{Style.RESET_ALL}")
            
        elif choice == "3":
            feedback = input("Provide feedback for re-extraction: ").strip()
            
            # Build refinement prompt and query the LLM again
            print(f"{Fore.YELLOW}Requesting re-extraction...{Style.RESET_ALL}")
            refinement_prompt = build_refinement_prompt(document_text, field, data, feedback)
            response = query_llm(refinement_prompt)
            
            if response:
                # Parse the refined response
                try:
                    cleaned_response = clean_json_response(response)
                    refined_data = json.loads(cleaned_response)
                    
                    # Extract the new value and confidence
                    if isinstance(refined_data, dict):
                        if "value" in refined_data:
                            new_value = refined_data["value"]
                            new_confidence = refined_data.get("confidence", 0.8)  # Default to 0.8 if not provided
                        else:
                            new_value = list(refined_data.values())[0]
                            new_confidence = 0.8
                    else:
                        new_value = refined_data
                        new_confidence = 0.8
                        
                    # Calculate external confidence
                    external_confidence = calculate_external_confidence(field, new_value, document_text)
                    combined_confidence = combine_confidences(new_confidence, external_confidence)
                    
                    # Update the result
                    updated_results[field]["value"] = new_value
                    updated_results[field]["confidence"] = combined_confidence
                    
                    print(f"{Fore.GREEN}Field re-extracted: \"{new_value}\" (Confidence: {combined_confidence:.1%}){Style.RESET_ALL}")
                except Exception as e:
                    print(f"{Fore.RED}Error processing refinement: {e}{Style.RESET_ALL}")
                    print("Keeping original value.")
            else:
                print(f"{Fore.RED}Re-extraction failed. Keeping original value.{Style.RESET_ALL}")
                
        elif choice == "4":
            print(f"{Fore.YELLOW}Skipped for now.{Style.RESET_ALL}")
            
        else:
            print(f"{Fore.RED}Invalid choice. Keeping original value.{Style.RESET_ALL}")
            
    print("\n" + "="*80)
    print(f"{Fore.GREEN}Review completed. Thank you for your feedback!{Style.RESET_ALL}")
    print("="*80)
    
    return updated_results

# --- FUNCTION TO UPDATE FIELD STATISTICS ---
def update_field_stats(results, stats_file=FIELD_STATS_FILE):
    """Track field extraction statistics over time"""
    # Initialize or load existing stats
    if os.path.exists(stats_file):
        try:
            stats_df = pd.read_csv(stats_file)
        except:
            # Create a new DataFrame if the file exists but can't be read
            stats_df = pd.DataFrame({
                'field': FIELDS,
                'extractions': 0,
                'found_count': 0,
                'not_found_count': 0,
                'avg_confidence': 0.0,
                'low_confidence_count': 0
            })
    else:
        # Create a new DataFrame if the file doesn't exist
        stats_df = pd.DataFrame({
            'field': FIELDS,
            'extractions': 0,
            'found_count': 0,
            'not_found_count': 0,
            'avg_confidence': 0.0,
            'low_confidence_count': 0
        })
        
    # Set field as index for easier updates
    stats_df.set_index('field', inplace=True)
    
    # Update stats for each field
    for field, data in results.items():
        if field in stats_df.index:
            # Increment extraction count
            stats_df.at[field, 'extractions'] += 1
            
            # Update found/not found counts
            if data['value'] == 'Not Found':
                stats_df.at[field, 'not_found_count'] += 1
            else:
                stats_df.at[field, 'found_count'] += 1
                
            # Update confidence stats
            current_avg = stats_df.at[field, 'avg_confidence']
            current_count = stats_df.at[field, 'extractions']
            new_count = current_count
            # Only include confidence for found values
            if data['value'] != 'Not Found':
                # Calculate new running average for confidence
                new_avg = ((current_avg * (current_count - 1)) + data['confidence']) / current_count
                stats_df.at[field, 'avg_confidence'] = new_avg
                
                # Update low confidence count
                if data['confidence'] < CONFIDENCE_THRESHOLD:
                    stats_df.at[field, 'low_confidence_count'] += 1
    
    # Reset index for saving
    stats_df = stats_df.reset_index()
    
    # Save updated stats
    stats_df.to_csv(stats_file, index=False)
    
    # Return summary of the stats
    return stats_df

# --- FUNCTION TO DISPLAY FIELD EXTRACTION DASHBOARD ---
def display_field_dashboard(stats_file=FIELD_STATS_FILE):
    """Display a simple dashboard of field extraction statistics"""
    if not os.path.exists(stats_file):
        print(f"{Fore.YELLOW}No extraction statistics available yet.{Style.RESET_ALL}")
        return
        
    try:
        # Load stats
        stats_df = pd.read_csv(stats_file)
        
        print("\n" + "="*80)
        print(f"{Fore.CYAN}FIELD EXTRACTION DASHBOARD{Style.RESET_ALL}")
        print("="*80)
        
        # Calculate success rate
        stats_df['success_rate'] = stats_df['found_count'] / stats_df['extractions'].where(stats_df['extractions'] > 0, 1)
        
        # Sort by success rate ascending (most problematic fields first)
        sorted_df = stats_df.sort_values('success_rate')
        
        print(f"{'FIELD':<20} {'EXTRACTIONS':^10} {'SUCCESS RATE':^15} {'AVG CONFIDENCE':^15} {'LOW CONF %':^15}")
        print("-"*80)
        
        for _, row in sorted_df.iterrows():
            field = row['field']
            extractions = row['extractions']
            
            if extractions > 0:
                success_rate = row['success_rate']
                avg_confidence = row['avg_confidence']
                low_conf_percent = row['low_confidence_count'] / extractions if extractions > 0 else 0
                
                # Color coding
                if success_rate >= 0.9:
                    success_color = Fore.GREEN
                elif success_rate >= 0.7:
                    success_color = Fore.YELLOW
                else:
                    success_color = Fore.RED
                    
                if avg_confidence >= 0.8:
                    conf_color = Fore.GREEN
                elif avg_confidence >= CONFIDENCE_THRESHOLD:
                    conf_color = Fore.YELLOW
                else:
                    conf_color = Fore.RED
                
                print(f"{field:<20} {extractions:^10} {success_color}{success_rate:^15.1%}{Style.RESET_ALL} "
                      f"{conf_color}{avg_confidence:^15.1%}{Style.RESET_ALL} {low_conf_percent:^15.1%}")
        
        print("-"*80)
        print(f"Total documents processed: {stats_df['extractions'].max()}")
        
        # Identify most problematic fields
        problem_fields = sorted_df[sorted_df['success_rate'] < 0.7]
        if not problem_fields.empty:
            print(f"\n{Fore.YELLOW}Fields with lowest success rates:{Style.RESET_ALL}")
            for _, row in problem_fields.iterrows():
                print(f"- {row['field']}: {row['success_rate']:.1%} success rate")
        
        print("="*80)
    except Exception as e:
        print(f"{Fore.RED}Error displaying dashboard: {e}{Style.RESET_ALL}")

# --- FUNCTION TO GET USER FEEDBACK FOR A FIELD ACROSS ALL DOCUMENTS (IMPROVED) ---
def get_field_feedback_across_documents(field, document_results, document_texts, filenames):
    """Get user feedback for a specific field and apply it across all documents - improved workflow"""
    print("\n" + "="*80)
    print(f"{Fore.CYAN}REVIEWING FIELD: {field} ACROSS ALL DOCUMENTS{Style.RESET_ALL}")
    print("="*80)
    
    # Get all values for this field across documents
    field_values = []
    
    for i, (result, filename, doc_text) in enumerate(zip(document_results, filenames, document_texts)):
        if field in result:
            field_data = result[field]
            confidence = field_data["confidence"]
            value = field_data["value"]
            
            if value != "Not Found":
                field_values.append({
                    "value": value,
                    "confidence": confidence,
                    "doc_index": i,
                    "filename": filename,
                    "doc_text": doc_text
                })
    
    # Show current values for this field across all documents
    print(f"\nCurrent values for {field}:")
    for i, (result, filename) in enumerate(zip(document_results, filenames)):
        field_data = result.get(field, {"value": "Not Found", "confidence": 0.0})
        value = field_data["value"]
        confidence = field_data["confidence"]
        
        # Color-code based on confidence
        if confidence >= 0.8:
            confidence_color = Fore.GREEN
        elif confidence >= CONFIDENCE_THRESHOLD:
            confidence_color = Fore.YELLOW
        else:
            confidence_color = Fore.RED
            
        confidence_str = f"{confidence:.1%}"
        print(f"[{i+1}] {filename}: {value} {confidence_color}(Confidence: {confidence_str}){Style.RESET_ALL}")
    
    updated_results = document_results.copy()
    
    # AUTO-SELECT THE BEST VALUE
    if len(field_values) > 0:
        # Sort values by confidence (highest first)
        field_values.sort(key=lambda x: x["confidence"], reverse=True)
        best_value = field_values[0]
        
        print(f"\n{Fore.CYAN}RECOMMENDED VALUE:{Style.RESET_ALL}")
        confidence_color = Fore.GREEN if best_value["confidence"] >= 0.8 else (Fore.YELLOW if best_value["confidence"] >= CONFIDENCE_THRESHOLD else Fore.RED)
        print(f"Value: \"{best_value['value']}\"")
        print(f"Source: {best_value['filename']}")
        print(f"Confidence: {confidence_color}{best_value['confidence']:.1%}{Style.RESET_ALL}")
        
        print("\nOptions:")
        print("1. Accept recommended value for all documents")
        print("2. Change the value manually")
        print("3. Request re-extraction with feedback for all documents")
        print("4. Skip for now")
        
        choice = input("\nYour choice (1-4): ").strip()
        
        if choice == "1" or choice == "":  # Default to accept if user just presses Enter
            print(f"{Fore.GREEN}Recommended value accepted.{Style.RESET_ALL}")
            # Apply the best value to all documents and set high confidence
            for result in updated_results:
                if field in result:
                    result[field]["value"] = best_value["value"]
                    result[field]["confidence"] = 1.0  # User-approved value gets 100% confidence
                else:
                    result[field] = {"value": best_value["value"], "confidence": 1.0}
                
        elif choice == "2":
            # Allow user to enter a new value
            new_value = input(f"Enter new value for {field}: ").strip()
            
            # Apply the new value to all documents
            for result in updated_results:
                if field in result:
                    result[field]["value"] = new_value
                    result[field]["confidence"] = 1.0  # Human-provided value gets 100% confidence
                else:
                    result[field] = {"value": new_value, "confidence": 1.0}
                    
            print(f"{Fore.GREEN}New value applied to all documents.{Style.RESET_ALL}")
                
        elif choice == "3":
            feedback = input("Provide feedback for re-extraction of this field: ").strip()
            
            print(f"{Fore.YELLOW}Requesting re-extraction for all documents...{Style.RESET_ALL}")
            
            # Apply re-extraction to all documents
            for i, (result, doc_text, filename) in enumerate(zip(updated_results, document_texts, filenames)):
                field_data = result.get(field, {"value": "Not Found", "confidence": 0.0})
                
                print(f"\nRe-extracting {field} for document [{i+1}] {filename}...")
                
                refinement_prompt = build_refinement_prompt(doc_text, field, field_data, feedback)
                response = query_llm(refinement_prompt)
                
                if response:
                    # Parse the refined response
                    try:
                        cleaned_response = clean_json_response(response)
                        refined_data = json.loads(cleaned_response)
                        
                        # Extract the new value and confidence
                        if isinstance(refined_data, dict):
                            if "value" in refined_data:
                                new_value = refined_data["value"]
                                new_confidence = refined_data.get("confidence", 0.8)  # Default to 0.8 if not provided
                            else:
                                new_value = list(refined_data.values())[0]
                                new_confidence = 0.8
                        else:
                            new_value = refined_data
                            new_confidence = 0.8
                            
                        # Calculate external confidence
                        external_confidence = calculate_external_confidence(field, new_value, doc_text)
                        combined_confidence = combine_confidences(new_confidence, external_confidence)
                        
                        # Update the result
                        if field in result:
                            result[field]["value"] = new_value
                            result[field]["confidence"] = combined_confidence
                        else:
                            result[field] = {"value": new_value, "confidence": combined_confidence}
                        
                        print(f"{Fore.GREEN}Field re-extracted: \"{new_value}\" (Confidence: {combined_confidence:.1%}){Style.RESET_ALL}")
                    except Exception as e:
                        print(f"{Fore.RED}Error processing refinement: {e}{Style.RESET_ALL}")
                        print("Keeping original value.")
                else:
                    print(f"{Fore.RED}Re-extraction failed. Keeping original value.{Style.RESET_ALL}")
                    
        elif choice == "4":
            print(f"{Fore.YELLOW}Skipped for now.{Style.RESET_ALL}")
            
        else:
            print(f"{Fore.RED}Invalid choice. Accepting recommended value.{Style.RESET_ALL}")
            # Default to accepting the recommended value
            for result in updated_results:
                if field in result:
                    result[field]["value"] = best_value["value"]
                    result[field]["confidence"] = 1.0
                else:
                    result[field] = {"value": best_value["value"], "confidence": 1.0}
    else:
        # No values found for this field
        print(f"\n{Fore.YELLOW}No values found in any document for field: {field}{Style.RESET_ALL}")
        
        # Ask if user wants to enter a value manually
        if input(f"Enter a value manually? (y/n): ").strip().lower() == 'y':
            new_value = input(f"Enter value for {field}: ").strip()
            
            # Apply to all documents
            for result in updated_results:
                if field in result:
                    result[field]["value"] = new_value
                    result[field]["confidence"] = 1.0
                else:
                    result[field] = {"value": new_value, "confidence": 1.0}
                    
            print(f"{Fore.GREEN}Manual value applied to all documents.{Style.RESET_ALL}")
        else:
            # Leave as "Not Found" for all documents
            for result in updated_results:
                if field in result:
                    result[field]["value"] = "Not Found"
                    result[field]["confidence"] = 0.0
                else:
                    result[field] = {"value": "Not Found", "confidence": 0.0}
                    
            print(f"{Fore.YELLOW}Field left as 'Not Found' for all documents.{Style.RESET_ALL}")
            
    print("\n" + "="*80)
    
    return updated_results

# --- FUNCTION TO CREATE COMBINED REPORT ---
def create_combined_report(document_results, filenames, output_path):
    """Create a combined report from all document results"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Combined Extraction Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(time.strftime('%d/%m/%Y'))
    
    doc.add_paragraph("\n")
    
    # For each field, show values across all documents
    for field in FIELDS:
        # Add section header for the field
        doc.add_heading(f"Field: {field}", level=1)
        
        # Create a table for this field
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Document"
        header_cells[1].text = "Value"
        header_cells[2].text = "Confidence"
        
        # Add data rows
        for filename, result in zip(filenames, document_results):
            field_data = result.get(field, {"value": "Not Found", "confidence": 0.0})
            value = field_data["value"]
            confidence = field_data["confidence"]
            confidence_str = f"{confidence:.1%}"
            
            # Add a row to the table
            row_cells = table.add_row().cells
            row_cells[0].text = filename
            row_cells[1].text = value
            row_cells[2].text = confidence_str
        
        doc.add_paragraph("\n")
    
    # Save the document
    doc.save(output_path)
    print(f"\nCombined report saved to: {output_path}")
    return True

# --- FUNCTION TO CREATE FINAL CONSOLIDATED TEMPLATE ---
def create_final_template(document_results, filenames, doc_texts, template_path, output_path):
    """Create the final, consolidated template with the best data from all documents"""
    try:
        print(f"\n{Fore.CYAN}CREATING FINAL CONSOLIDATED REPORT{Style.RESET_ALL}")
        print("="*80)
        print("Analyzing all extracted data to find the best values for each field...")
        
        # Create a consolidated dataset that represents the best data from all documents
        final_data = {}
        field_sources = {}
        
        for field in FIELDS:
            # Get all values for this field across documents
            field_values = []
            
            for i, (result, filename, doc_text) in enumerate(zip(document_results, filenames, doc_texts)):
                if field in result:
                    field_data = result[field]
                    confidence = field_data["confidence"]
                    value = field_data["value"]
                    
                    if value != "Not Found":
                        field_values.append({
                            "value": value,
                            "confidence": confidence,
                            "doc_index": i,
                            "filename": filename,
                            "doc_text": doc_text
                        })
            
            # Sort values by confidence (highest first)
            field_values.sort(key=lambda x: x["confidence"], reverse=True)
            
            # Let the user select the best value if there are multiple options
            if len(field_values) > 0:
                print(f"\nField: {Fore.CYAN}{field}{Style.RESET_ALL}")
                
                # If there's only one value or the highest confidence is very high, auto-select it
                if len(field_values) == 1 or field_values[0]["confidence"] > 0.9:
                    best_value = field_values[0]
                    print(f"Auto-selected value (highest confidence): \"{best_value['value']}\" from {best_value['filename']} (Confidence: {best_value['confidence']:.1%})")
                    
                    # Store the value and its source
                    final_data[field] = {
                        "value": best_value["value"],
                        "confidence": best_value["confidence"]
                    }
                    field_sources[field] = best_value["filename"]
                else:
                    # Show options to the user
                    print("Multiple values found. Please select the best one:")
                    for i, val in enumerate(field_values):
                        confidence_color = Fore.GREEN if val["confidence"] >= 0.8 else (Fore.YELLOW if val["confidence"] >= CONFIDENCE_THRESHOLD else Fore.RED)
                        print(f"[{i+1}] \"{val['value']}\" from {val['filename']} {confidence_color}(Confidence: {val['confidence']:.1%}){Style.RESET_ALL}")
                    
                    print(f"[{len(field_values)+1}] Enter a new value manually")
                    print(f"[{len(field_values)+2}] Skip this field (leave empty)")
                    
                    # Get user selection
                    valid_selection = False
                    while not valid_selection:
                        try:
                            selection = input(f"Your choice (1-{len(field_values)+2}): ").strip()
                            
                            # Check if the user just pressed enter (default to highest confidence)
                            if selection == "":
                                selection = "1"
                                
                            selection = int(selection)
                            
                            if 1 <= selection <= len(field_values):
                                # User selected an existing value
                                selected_value = field_values[selection-1]
                                final_data[field] = {
                                    "value": selected_value["value"],
                                    "confidence": selected_value["confidence"]
                                }
                                field_sources[field] = selected_value["filename"]
                                print(f"{Fore.GREEN}Selected: \"{selected_value['value']}\" (Confidence: {selected_value['confidence']:.1%}){Style.RESET_ALL}")
                                valid_selection = True
                                
                            elif selection == len(field_values)+1:
                                # User wants to enter a new value manually
                                new_value = input(f"Enter new value for {field}: ").strip()
                                final_data[field] = {
                                    "value": new_value,
                                    "confidence": 1.0  # User-provided value gets 100% confidence
                                }
                                field_sources[field] = "Manual Entry"
                                print(f"{Fore.GREEN}Manual value saved: \"{new_value}\"{Style.RESET_ALL}")
                                valid_selection = True
                                
                            elif selection == len(field_values)+2:
                                # User wants to skip this field
                                final_data[field] = {
                                    "value": "",
                                    "confidence": 0.0
                                }
                                field_sources[field] = "Skipped"
                                print(f"{Fore.YELLOW}Field skipped.{Style.RESET_ALL}")
                                valid_selection = True
                                
                            else:
                                print(f"{Fore.RED}Invalid choice. Please try again.{Style.RESET_ALL}")
                                
                        except ValueError:
                            print(f"{Fore.RED}Please enter a number.{Style.RESET_ALL}")
            else:
                # No values found for this field
                print(f"\nField: {Fore.CYAN}{field}{Style.RESET_ALL}")
                print(f"{Fore.YELLOW}No values found in any document.{Style.RESET_ALL}")
                
                # Ask if user wants to enter a value manually
                if input(f"Enter a value manually? (y/n): ").strip().lower() == 'y':
                    new_value = input(f"Enter value for {field}: ").strip()
                    final_data[field] = {
                        "value": new_value,
                        "confidence": 1.0
                    }
                    field_sources[field] = "Manual Entry"
                    print(f"{Fore.GREEN}Manual value saved: \"{new_value}\"{Style.RESET_ALL}")
                else:
                    final_data[field] = {
                        "value": "",
                        "confidence": 0.0
                    }
                    field_sources[field] = "Empty"
                    print(f"{Fore.YELLOW}Field left empty.{Style.RESET_ALL}")
        
        # Fill the template with the final consolidated data
        print("\nFilling final template with consolidated data...")
        success = fill_template_improved(final_data, template_path, output_path)
        
        if success:
            print(f"\n{Fore.GREEN}Final template created successfully: {output_path}{Style.RESET_ALL}")
            
            # Create a field sources report
            sources_report_path = os.path.join(os.path.dirname(output_path), "Final_Template_Sources.txt")
            with open(sources_report_path, 'w', encoding='utf-8') as f:
                f.write("FINAL TEMPLATE SOURCES\n")
                f.write("=====================\n\n")
                f.write("This report shows the source for each field in the final template.\n\n")
                
                for field in FIELDS:
                    field_data = final_data.get(field, {"value": "", "confidence": 0.0})
                    value = field_data["value"]
                    confidence = field_data["confidence"]
                    source = field_sources.get(field, "Unknown")
                    
                    f.write(f"Field: {field}\n")
                    f.write(f"Value: {value}\n")
                    f.write(f"Confidence: {confidence:.1%}\n")
                    f.write(f"Source: {source}\n")
                    f.write("-" * 50 + "\n\n")
                    
            print(f"Field sources report created: {sources_report_path}")
            
            return True
        else:
            print(f"{Fore.RED}Failed to create final template.{Style.RESET_ALL}")
            return False
            
    except Exception as e:
        print(f"{Fore.RED}Error creating final template: {e}{Style.RESET_ALL}")
        return False

# --- MAIN EXTRACTION PIPELINE ---
def main():
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)

    # Read all documents at once
    print(f"\n{Fore.CYAN}Reading all documents from {FOLDER_PATH}...{Style.RESET_ALL}")
    documents, filenames = read_pdfs_from_folder(FOLDER_PATH)
    print(f"Found {len(documents)} documents.")

    # Process all documents for initial extraction before user review
    all_document_results = []
    doc_texts = []

    # Track stats across all documents
    all_extraction_stats = []

    # STEP 1: Process all documents in batch for initial extraction
    for i, (doc_text, filename) in enumerate(zip(documents, filenames)):
        print(f"\n{Fore.CYAN}[{i+1}/{len(documents)}] Initial extraction for: {filename}{Style.RESET_ALL}")
        
        # Initialize result dictionary
        result = {}
        
        # Process fields in batches
        for batch in batch_fields(FIELDS, BATCH_SIZE):
            # First extraction
            prompt = build_prompt(doc_text, batch, previous_values=result)
            response = query_llm(prompt)
            batch_values = parse_response(response, batch)
            
            # Enhance with external confidence calculation
            for field, data in batch_values.items():
                # Calculate external confidence
                external_confidence = calculate_external_confidence(field, data["value"], doc_text)
                
                # Combine confidences
                llm_confidence = data["confidence"]
                combined_confidence = combine_confidences(llm_confidence, external_confidence)
                
                # Update confidence
                batch_values[field]["confidence"] = combined_confidence
            
            # Update results
            result.update(batch_values)
            
            # Display interim results
            print_extraction_results(batch_values)
            
            # Add a small delay between batches
            time.sleep(1)
        
        # Save the document text and results for later review
        all_document_results.append(result)
        doc_texts.append(doc_text)

    # STEP 2: Now identify fields that need review across ALL documents
    print(f"\n{Fore.GREEN}Initial extraction of all documents complete. Starting field-by-field review phase...{Style.RESET_ALL}")
    
    # Collect fields that need review across all documents
    fields_to_review = set()
    for result in all_document_results:
        for field, data in result.items():
            confidence = data["confidence"]
            value = data["value"]
            
            # Flag fields with low confidence or "Not Found" values for review
            if confidence < CONFIDENCE_THRESHOLD or value == "Not Found":
                fields_to_review.add(field)
    
    # Sort fields to review based on their order in FIELDS list
    fields_to_review = sorted(fields_to_review, key=lambda x: FIELDS.index(x) if x in FIELDS else 999)
    
    print(f"\n{Fore.YELLOW}Fields requiring review across all documents:{Style.RESET_ALL}")
    for field in fields_to_review:
        print(f"- {field}")
    
    # Interactive review by field
    for field in fields_to_review:
        # Get feedback and updates for this field across all documents
        all_document_results = get_field_feedback_across_documents(field, all_document_results, doc_texts, filenames)
    
    print(f"\n{Fore.GREEN}Field review completed for all documents.{Style.RESET_ALL}")
    
    # Save the final results for each document
    for result in all_document_results:
        all_extraction_stats.append(result)
    
    # Create a combined report for all documents
    combined_report_path = os.path.join(OUTPUT_FOLDER, "Combined_Extraction_Report.docx")
    create_combined_report(all_document_results, filenames, combined_report_path)
    
    # Create the final template with user selection for each field
    template_path = TEMPLATE_DOCX
    final_template_path = os.path.join(OUTPUT_FOLDER, "Template_Final.docx")
    create_final_template(all_document_results, filenames, doc_texts, template_path, final_template_path)
    
    # Update field statistics
    for result in all_extraction_stats:
        update_field_stats(result)
    
    # Display the extraction dashboard
    display_field_dashboard()
    
    print(f"\n{Fore.GREEN}Processing complete. Check the {OUTPUT_FOLDER} directory for outputs:{Style.RESET_ALL}")
    print(f"1. Combined extraction report: {combined_report_path}")
    print(f"2. Final template: {final_template_path}")
    print(f"3. Field sources report: {os.path.join(OUTPUT_FOLDER, 'Final_Template_Sources.txt')}")
    print(f"4. Field extraction dashboard: {os.path.join(OUTPUT_FOLDER, 'field_extraction_dashboard.html')}")

# --- FIELD TRACKER DASHBOARD FUNCTION ---
def field_tracker_dashboard():
    """Generate a comprehensive field tracking dashboard as CSV and HTML"""
    if not os.path.exists(FIELD_STATS_FILE):
        print(f"{Fore.YELLOW}No extraction statistics available yet. Run the main extraction first.{Style.RESET_ALL}")
        return
    
    try:
        # Load stats
        stats_df = pd.read_csv(FIELD_STATS_FILE)
        
        # Calculate additional metrics
        stats_df['success_rate'] = stats_df['found_count'] / stats_df['extractions'].where(stats_df['extractions'] > 0, 1)
        stats_df['low_confidence_rate'] = stats_df['low_confidence_count'] / stats_df['extractions'].where(stats_df['extractions'] > 0, 1)
        
        # Sort by success rate
        sorted_df = stats_df.sort_values('success_rate')
        
        # Save enhanced stats to CSV
        enhanced_stats_file = os.path.join(OUTPUT_FOLDER, "field_extraction_dashboard.csv")
        sorted_df.to_csv(enhanced_stats_file, index=False)
        
        # Generate a simple HTML dashboard
        html_file = os.path.join(OUTPUT_FOLDER, "field_extraction_dashboard.html")
        
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Field Extraction Dashboard</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                h1 { color: #2c3e50; }
                table { border-collapse: collapse; width: 100%; margin-top: 20px; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: center; }
                th { background-color: #f2f2f2; }
                tr:nth-child(even) { background-color: #f9f9f9; }
                .good { background-color: #d4edda; }
                .medium { background-color: #fff3cd; }
                .poor { background-color: #f8d7da; }
                .summary { margin-top: 30px; }
            </style>
        </head>
        <body>
            <h1>Field Extraction Dashboard</h1>
            <p>Total documents processed: """ + str(stats_df['extractions'].max()) + """</p>
            
            <table>
                <tr>
                    <th>Field</th>
                    <th>Extractions</th>
                    <th>Success Rate</th>
                    <th>Avg Confidence</th>
                    <th>Low Confidence %</th>
                    <th>Status</th>
                </tr>
        """
        
        for _, row in sorted_df.iterrows():
            field = row['field']
            extractions = row['extractions']
            
            if extractions > 0:
                success_rate = row['success_rate']
                avg_confidence = row['avg_confidence']
                low_conf_percent = row['low_confidence_rate']
                
                # Status class
                if success_rate >= 0.9 and avg_confidence >= 0.8:
                    status_class = "good"
                    status_text = "Good"
                elif success_rate >= 0.7 and avg_confidence >= CONFIDENCE_THRESHOLD:
                    status_class = "medium"
                    status_text = "Needs Improvement"
                else:
                    status_class = "poor"
                    status_text = "Problematic"
                
                html_content += f"""
                <tr class="{status_class}">
                    <td>{field}</td>
                    <td>{extractions}</td>
                    <td>{success_rate:.1%}</td>
                    <td>{avg_confidence:.1%}</td>
                    <td>{low_conf_percent:.1%}</td>
                    <td>{status_text}</td>
                </tr>
                """
        
        # Add summary and recommendations
        problem_fields = sorted_df[sorted_df['success_rate'] < 0.7]
        html_content += """
            </table>
            
            <div class="summary">
                <h2>Summary and Recommendations</h2>
        """
        
        if not problem_fields.empty:
            html_content += "<h3>Problematic Fields:</h3><ul>"
            for _, row in problem_fields.iterrows():
                html_content += f"<li><strong>{row['field']}</strong>: {row['success_rate']:.1%} success rate</li>"
            html_content += "</ul>"
            
            html_content += """
                <h3>Recommendations:</h3>
                <ol>
                    <li>Review document formatting for problematic fields</li>
                    <li>Consider adjusting the extraction prompt for these specific fields</li>
                    <li>Add more training examples for the problematic fields</li>
                </ol>
            """
        else:
            html_content += "<p>All fields are performing well. Continue monitoring.</p>"
        
        html_content += """
            </div>
        </body>
        </html>
        """
        
        # Write the HTML file
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        print(f"\n{Fore.GREEN}Dashboard generated successfully:{Style.RESET_ALL}")
        print(f"- CSV: {enhanced_stats_file}")
        print(f"- HTML: {html_file}")
        
    except Exception as e:
        print(f"{Fore.RED}Error generating dashboard: {e}{Style.RESET_ALL}")

# Entry point
if __name__ == "__main__":
    # Run the main extraction pipeline
    main()
    
    # Generate the field tracker dashboard
    field_tracker_dashboard()