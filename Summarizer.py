import os
import time
import json
import fitz  # PyMuPDF
import re
import urllib3
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- SETTINGS ---
FOLDER_PATH = "Gynecology Anna Maria Fortuny Pla"
OUTPUT_FOLDER = "Extracted_Documents"
FINAL_DOCX = "Final_Extracted_Report.docx"
TEMPLATE_DOCX = "Template.docx"
FILLED_TEMPLATE = "Filled_Template.docx"
API_TOKEN = "YOUR_ACCESS_TOKEN"  # Replace with your API token

# --- FIELD DEFINITIONS ---
FIELDS = [
    "DATE", "INCIDENT_DATE", "PATIENT_NAME", "DOB", "MEDICAL_DOCTOR", "LAWYER",
    "SPECIALITY", "CENTER", "JUDGE", "REFERENCE", "SUMMARY", "PRACTICE_DETAILS",
    "FORENSIC_REPORT", "EVALUATION_UML", "DIAGNOSTIC_CODE", "PROCEDURE_CODE",
    "SEQUELS", "DEFINITIVE?", "QUANTUM", "RISK_ASSESSMENT"
]
BATCH_SIZE = 3

# --- FUNCTION TO CALL THE API ---
def call_model(prompt):
    urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
    destination_url = "https://grupmedai-api-des.itcomb.cat/llm/generate"
    headers = {
        'Authorization':  f'Bearer {API_TOKEN}',
        'Content-Type': 'application/json'
    }
    data = {
        "model": "llama3.1:8b",
        "prompt": prompt 
    }
    try:
        response = requests.post(destination_url, headers=headers, data=json.dumps(data), verify=False)
        if response.status_code == 200:
            response_data = response.json().get('text', '')
            return response_data
        else:
            print(f"API error: Status code {response.status_code}")
            return False
    except Exception as e:
        print(f"Request error: {e}")
        return False

# --- FUNCTION TO BATCH FIELDS ---
def batch_fields(field_list, batch_size):
    for i in range(0, len(field_list), batch_size):
        yield field_list[i:i+batch_size]

# --- FUNCTION TO READ PDFs ---
def read_pdfs_from_folder(folder_path):
    documents = []
    filenames = []
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".pdf"):
            file_path = os.path.join(folder_path, filename)
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            documents.append(text)
            filenames.append(os.path.splitext(filename)[0])
    return documents, filenames

# --- FUNCTION TO BUILD PROMPT ---
def build_prompt(document_text, fields_batch):
    fields_prompt = "\n".join([f"- {field}" for field in fields_batch])
    prompt = f"""
You are an AI assistant extracting key information from a medico-legal report.

Instructions:
- For each field, extract only the direct, specific answer.
- For PERSON fields (like PATIENT_NAME, MEDICAL_DOCTOR), extract only the full name. No age, specialty, or biography.
- For DATE fields, extract only the date in format DD/MM/YYYY.
- Be concise. No paragraphs, no explanations, no extra text.
- If the field is missing in the document, respond with "Not Found".
- YOU MUST RESPOND WITH VALID JSON FORMAT ONLY. 
- Format your response as: {{"FIELD1": "value1", "FIELD2": "value2", ...}}
- Do not include any additional text outside of the JSON object.

Fields to extract:
{fields_prompt}

Document Text:
{document_text}
"""
    return prompt

# --- FUNCTION TO QUERY LLM ---
def query_llm(prompt):
    try:
        response = call_model(prompt)
        return response
    except Exception as e:
        print(f"Error querying LLM: {e}")
        return None

# --- FUNCTION TO CLEAN LLM RESPONSE ---
def clean_json_response(response_text):
    if response_text is None or response_text is False:
        return None
        
    # Find JSON object in response if it's embedded in other text
    json_pattern = r'({.*})'
    match = re.search(json_pattern, response_text, re.DOTALL)
    
    if match:
        potential_json = match.group(1)
        # Clean up common JSON formatting issues
        potential_json = potential_json.replace("'", '"')  # Replace single quotes with double quotes
        potential_json = re.sub(r',\s*}', '}', potential_json)  # Remove trailing commas
        
        return potential_json
    else:
        return response_text  # Return original if no JSON-like pattern found

# --- FUNCTION TO PARSE RESPONSE ---
def parse_response(response_text, fields_batch):
    field_values = {}
    
    # Initialize with "Not Found" values
    for field in fields_batch:
        field_values[field] = "Not Found"
    
    if response_text is None or response_text is False:
        return field_values
    
    # Clean the response to get valid JSON
    cleaned_response = clean_json_response(response_text)
    
    try:
        data = json.loads(cleaned_response)
        for field in fields_batch:
            value = data.get(field, "Not Found")
            if isinstance(value, dict) and 'Name' in value:
                value = value['Name']
            field_values[field] = value
    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
        print(f"Response text: {response_text}")
        print(f"Cleaned response: {cleaned_response}")
        
        # Fallback: Try to extract field-value pairs directly from text using regex
        for field in fields_batch:
            pattern = rf'"{field}"\s*:\s*"([^"]*)"'
            match = re.search(pattern, cleaned_response)
            if match:
                field_values[field] = match.group(1)
    
    return field_values

# --- FUNCTION TO CREATE FINAL DOCX ---
def create_final_docx(data, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for field in FIELDS:
        value = data.get(field, "Not Found")
        # Convert any non-string values to strings
        if not isinstance(value, str):
            value = str(value)
        doc.add_paragraph(f"{field}: {value}")

    doc.save(output_path)

# --- FUNCTION TO ENSURE STRING VALUES ---
def ensure_string_values(data_dict):
    """Convert all values in a dictionary to strings to prevent template filling errors."""
    for key, value in data_dict.items():
        if not isinstance(value, str):
            data_dict[key] = str(value)
    return data_dict

# --- FUNCTION TO FILL TEMPLATE ---
def fill_template(data, template_path, output_path):
    try:
        # Convert all data values to strings to prevent replace() errors
        data = ensure_string_values(data)
        
        # Load the template document
        doc = Document(template_path)
        
        # Process each table in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Check each paragraph in the cell
                    for paragraph in cell.paragraphs:
                        # Find all placeholders in the format {FIELD_NAME}
                        placeholder_pattern = r'\{([A-Z_?]+)\}'
                        matches = re.findall(placeholder_pattern, paragraph.text)
                        
                        # Replace placeholders with extracted values
                        for field in matches:
                            if field in data:
                                value = data.get(field, "Not Found")
                                paragraph.text = paragraph.text.replace(f"{{{field}}}", value)
        
        # Process paragraphs outside of tables
        for paragraph in doc.paragraphs:
            placeholder_pattern = r'\{([A-Z_?]+)\}'
            matches = re.findall(placeholder_pattern, paragraph.text)
            
            for field in matches:
                if field in data:
                    value = data.get(field, "Not Found")
                    paragraph.text = paragraph.text.replace(f"{{{field}}}", value)
        
        # Save the filled template
        doc.save(output_path)
        print(f"Template filled and saved to: {output_path}")
        return True
    except Exception as e:
        print(f"Error filling template: {e}")
        return False

# --- IMPROVED FUNCTION TO FILL TEMPLATE ---
def fill_template_improved(data, template_path, output_path):
    """
    Improved template filling that better handles complex documents and run-level formatting.
    """
    try:
        # Convert all data values to strings
        data = ensure_string_values(data)
        
        # Load the template document
        doc = Document(template_path)
        
        # Helper function to replace text while preserving formatting
        def replace_placeholder_in_paragraph(paragraph, field, value):
            # Get placeholder text pattern
            placeholder = f"{{{field}}}"
            
            # Skip if placeholder not in paragraph
            if placeholder not in paragraph.text:
                return False
                
            # Identify which run contains the placeholder
            found = False
            runs_with_placeholder = []
            placeholder_start_run = None
            placeholder_text_found = ""
            
            # First find all runs containing parts of the placeholder
            for i, run in enumerate(paragraph.runs):
                if placeholder in run.text:
                    # Simple case: entire placeholder in one run
                    run.text = run.text.replace(placeholder, value)
                    found = True
                    break
                    
                # Check for partial placeholder (placeholder split across runs)
                if placeholder_start_run is None and "{" in run.text and run.text.endswith("{"):
                    # Start of a potential placeholder
                    placeholder_start_run = i
                    placeholder_text_found = run.text[run.text.index("{"):]
                elif placeholder_start_run is not None:
                    # Continue building the placeholder
                    placeholder_text_found += run.text
                    
                    # Check if we've found the complete placeholder
                    if placeholder in placeholder_text_found:
                        runs_with_placeholder = list(range(placeholder_start_run, i + 1))
                        break
                    
                    # Check if this run closes a placeholder but not our target
                    if "}" in run.text:
                        placeholder_start_run = None
                        placeholder_text_found = ""
            
            # Handle multi-run placeholders
            if not found and runs_with_placeholder:
                # First run gets everything up to the placeholder start + the replacement value
                start_run = paragraph.runs[runs_with_placeholder[0]]
                placeholder_start_pos = start_run.text.find("{")
                start_run.text = start_run.text[:placeholder_start_pos] + value
                
                # Delete text from middle runs
                for i in runs_with_placeholder[1:-1]:
                    paragraph.runs[i].text = ""
                
                # Last run gets everything after the placeholder end
                end_run = paragraph.runs[runs_with_placeholder[-1]]
                placeholder_end_pos = end_run.text.find("}") + 1
                if placeholder_end_pos < len(end_run.text):
                    end_run.text = end_run.text[placeholder_end_pos:]
                else:
                    end_run.text = ""
                    
                found = True
                
            return found
        
        # Process table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for field, value in data.items():
                            replace_placeholder_in_paragraph(paragraph, field, value)
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            for field, value in data.items():
                replace_placeholder_in_paragraph(paragraph, field, value)
                
        # Save the filled template
        doc.save(output_path)
        print(f"Template filled and saved to: {output_path}")
        return True
    except Exception as e:
        print(f"Error filling template (improved method): {e}")
        return False

# --- FUNCTION TO PRINT EXTRACTION RESULTS ---
def print_extraction_results(result):
    print("\nExtraction Results:")
    print("-" * 40)
    for field, value in result.items():
        if value != "Not Found":
            print(f"{field}: {value}")
    print("-" * 40)

# --- MAIN EXTRACTION PIPELINE ---
def main():
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)

    documents, filenames = read_pdfs_from_folder(FOLDER_PATH)

    final_result = {field: "Not Found" for field in FIELDS}

    for doc_text, filename in zip(documents, filenames):
        print(f"Processing: {filename}")
        result = {}
        for batch in batch_fields(FIELDS, BATCH_SIZE):
            prompt = build_prompt(doc_text, batch)
            response = query_llm(prompt)
            batch_values = parse_response(response, batch)
            result.update(batch_values)
            print_extraction_results(batch_values)  # Print results for debugging
            time.sleep(2)

        # Merge fields: only overwrite if final_result still has "Not Found"
        for field, value in result.items():
            if final_result[field] == "Not Found" and value != "Not Found":
                final_result[field] = value

    # Save extracted data to regular docx report
    output_path = os.path.join(OUTPUT_FOLDER, FINAL_DOCX)
    create_final_docx(final_result, output_path)
    print(f"Final report saved to: {output_path}")
    
    # Try filling the template with all methods
    template_path = TEMPLATE_DOCX  # Use the template from your file upload
    
    # Standard method
    standard_filled_path = os.path.join(OUTPUT_FOLDER, "Standard_" + FILLED_TEMPLATE)
    standard_success = fill_template(final_result, template_path, standard_filled_path)
    
    # Improved method
    improved_filled_path = os.path.join(OUTPUT_FOLDER, "Improved_" + FILLED_TEMPLATE)
    improved_success = fill_template_improved(final_result, template_path, improved_filled_path)
    
    # Let the user know which filled template to use
    if standard_success and improved_success:
        print(f"Both template filling methods succeeded. Compare the outputs to see which looks better:")
        print(f"1. Standard: {standard_filled_path}")
        print(f"2. Improved: {improved_filled_path}")
    elif standard_success:
        print(f"Template successfully filled using standard method: {standard_filled_path}")
    elif improved_success:
        print(f"Template successfully filled using improved method: {improved_filled_path}")
    else:
        print("All template filling methods failed. Check the error messages above.")

# Alternative main function that uses debug print statements
def debug_main():
    """Version of main with additional debugging for template filling"""
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)

    # For debugging, use a simple test dictionary
    test_data = {
        "DATE": "01/01/2025",
        "INCIDENT_DATE": "18/05/2013",
        "PATIENT_NAME": "Ana Garcia Lopez",
        "DOB": "15/03/1943",
        "MEDICAL_DOCTOR": "Dr. Alejandro Sánchez Gómez",
        "LAWYER": "Marta Gomez",
        "SPECIALITY": "General surgery",
        "CENTER": "Hospital Residència Sant Camil",
        "JUDGE": "Not Found",
        "REFERENCE": "179/2019",
        "SUMMARY": "Patient underwent exploratory laparotomy resulting in medical treatment side effects",
        "PRACTICE_DETAILS": "Exploratory laparotomy performed to address arterial disorder",
        "FORENSIC_REPORT": "Not Found",
        "EVALUATION_UML": "Not Found",
        "DIAGNOSTIC_CODE": "I779",
        "PROCEDURE_CODE": "54.11",
        "SEQUELS": "57",
        "DEFINITIVE?": "Not Found",
        "QUANTUM": "Not Found",
        "RISK_ASSESSMENT": "High due to lack of informed consent"
    }

    # Print the test data
    print("Test data for template filling:")
    for k, v in test_data.items():
        print(f"{k}: {v} (type: {type(v)})")
    
    # Try filling the template with test data
    template_path = TEMPLATE_DOCX
    
    # Make sure all values are strings
    test_data = ensure_string_values(test_data)
    
    # Try standard method
    test_standard_path = os.path.join(OUTPUT_FOLDER, "Test_Standard_" + FILLED_TEMPLATE)
    print(f"\nTrying standard method to fill template...")
    standard_success = fill_template(test_data, template_path, test_standard_path)
    
    # Try improved method
    test_improved_path = os.path.join(OUTPUT_FOLDER, "Test_Improved_" + FILLED_TEMPLATE)
    print(f"\nTrying improved method to fill template...")
    improved_success = fill_template_improved(test_data, template_path, test_improved_path)
    
    # Results
    if standard_success or improved_success:
        print("\nAt least one template filling method worked with test data.")
        print("Now trying the regular extraction process...")
        main()
    else:
        print("\nBoth template filling methods failed with test data.")
        print("There may be an issue with the template format or structure.")

if __name__ == "__main__":
    # Uncomment the line below to run debugging mode
    # debug_main()
    
    # Normal execution
    main()