from docx import Document
import datetime

def generate_report(data, template_path='UML Plantilla Kopie.docx', output_path='Generated_Report.docx'):
    """
    Generates a Word report from a template by replacing placeholders with given data.
    
    :param data: Dictionary containing replacement values for placeholders
    :param template_path: Path to the template document
    :param output_path: Path to save the generated document
    """
    # Load the template document
    doc = Document(template_path)
    
    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        for key, value in data.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    
    # Save the modified document
    doc.save(output_path)
    print(f"Report generated: {output_path}")

# Example usage
if __name__ == "__main__":
    custom_data = {
        "{DATE}": datetime.date.today().strftime("%d/%m/%Y"),
        "{INCIDENT_DATE}": "06/07/2019",
        "{PATIENT_NAME}": "John Doe",
        "{DOB}": "30/06/1976",
        "{MEDICAL_DOCTOR}": "Dr. Alice Smith",
        "{LAWYER}": "US",
        "{SPECIALITY}": "Urology",
        "{CENTER}": "General Hospital",
        "{JUDGE}": "Extrajudicial",
        "{SUMMARY}": "Diagnostic error in testicular torsion with final orchiectomy.",
        "{PRACTICE_DETAILS}": "Patient experienced pain for 3 days before seeking care.",
        "{DIAGNOSTIC_CODE}": "CIM-9 MC: 12345",
        "{PROCEDURE_CODE}": "CIM-9 MC: 67890",
        "{RISK_ASSESSMENT}": "Defend - Lack of Damage",
    }
    
    generate_report(custom_data)
